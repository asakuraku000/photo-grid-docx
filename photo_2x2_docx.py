"""
A4 Photo Grid → DOCX Generator
Layouts: 1x1 (full page), 1x2, 1x3, 2x2, 3x3, 4x4
- Add images via file dialog or paste path
- Responsive left-to-right card grid (auto-columns on resize)
- Drag to reorder before generating
- Progress bar while building the DOCX
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import threading
import math
from io import BytesIO

# ── Pillow (optional, for thumbnails + image resizing) ───────────────────────
try:
    from PIL import Image, ImageTk
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ── python-docx ──────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────────────────────────
#  Layout definitions
# ─────────────────────────────────────────────────────────────────────────────

LAYOUTS = {
    "1×1  – Full page":   (1, 1),
    "1×2  – 2 portrait":  (1, 2),
    "1×3  – 3 portrait":  (1, 3),
    "2×2  – 4 per page":  (2, 2),
    "2×3  – 6 per page":  (2, 3),
    "3×3  – 9 per page":  (3, 3),
    "4×4  – 16 per page": (4, 4),
}
DEFAULT_LAYOUT = "2×2  – 4 per page"

A4_W_MM   = 210
A4_H_MM   = 297
MARGIN_MM = 10
GAP_MM    = 3

# DPI used when resizing images before embedding.
EMBED_DPI = 200

# ── Card / thumbnail sizing ───────────────────────────────────────────────────
THUMB_SIZE = (110, 110)   # square thumbnails in the responsive card grid


def cell_dims(cols: int, rows: int):
    """Return (cell_w_mm, cell_h_mm) for a given grid on A4."""
    usable_w = A4_W_MM - MARGIN_MM * 2 - GAP_MM * (cols - 1)
    usable_h = A4_H_MM - MARGIN_MM * 2 - GAP_MM * (rows - 1)
    return usable_w / cols, usable_h / rows


# ─────────────────────────────────────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────────────────────────────────────

def make_thumbnail(path: str):
    if not HAS_PIL:
        return None
    try:
        img = Image.open(path)
        img.thumbnail(THUMB_SIZE, Image.LANCZOS)
        bg = Image.new("RGB", THUMB_SIZE, (45, 45, 45))
        ox = (THUMB_SIZE[0] - img.width)  // 2
        oy = (THUMB_SIZE[1] - img.height) // 2
        bg.paste(img, (ox, oy))
        return ImageTk.PhotoImage(bg)
    except Exception:
        return None


def resize_for_embed(img_path: str, cell_w_mm: float, cell_h_mm: float):
    """
    Resize + CENTER-CROP the image so it fills the cell EXACTLY (cover-fit),
    at EMBED_DPI, then return a JPEG BytesIO buffer.
    Returns None if PIL is unavailable or on any error (caller falls back).
    """
    if not HAS_PIL:
        return None
    try:
        target_w = max(1, int(round(cell_w_mm / 25.4 * EMBED_DPI)))
        target_h = max(1, int(round(cell_h_mm / 25.4 * EMBED_DPI)))

        with Image.open(img_path) as img:
            try:
                from PIL import ImageOps
                img = ImageOps.exif_transpose(img)
            except Exception:
                pass

            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            src_w, src_h = img.size
            scale  = max(target_w / src_w, target_h / src_h)
            new_w  = max(target_w, int(math.ceil(src_w * scale)))
            new_h  = max(target_h, int(math.ceil(src_h * scale)))
            img    = img.resize((new_w, new_h), Image.LANCZOS)

            left = (new_w - target_w) // 2
            top  = (new_h - target_h) // 2
            img  = img.crop((left, top, left + target_w, top + target_h))

            buf = BytesIO()
            img.save(buf, format="JPEG", quality=85, optimize=True)
            buf.seek(0)
            return buf
    except Exception:
        return None


def install_missing():
    import subprocess, sys
    missing = []
    if not HAS_PIL:   missing.append("Pillow")
    if not HAS_DOCX:  missing.append("python-docx")
    if not missing:
        return ""
    subprocess.check_call([sys.executable, "-m", "pip", "install"] + missing,
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return "Installed: " + ", ".join(missing)


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX builder
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start),
                      ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="FFFFFF", size=0):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"),   "single")
        node.set(qn("w:sz"),    str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        tblBorders.append(node)
    tblPr.append(tblBorders)


def set_row_height(row, height_mm):
    tr   = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = OxmlElement("w:trHeight")
    twips = int(height_mm * 56.69)
    trHeight.set(qn("w:val"),   str(twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def set_row_cant_split(row):
    """Prevent a table row from being split across two pages."""
    tr   = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trPr.append(OxmlElement("w:cantSplit"))


def build_docx(image_paths: list, output_path: str,
               cols: int = 2, rows: int = 2,
               progress_cb=None, done_cb=None):
    """
    Build a DOCX with images in a cols×rows A4 grid.
    Images are resized to cell dimensions before embedding so the file stays
    small and memory usage stays low even for 40+ photos.
    """
    if not HAS_DOCX:
        if done_cb:
            done_cb("python-docx not installed. Run: pip install python-docx")
        return

    try:
        doc     = Document()
        section = doc.sections[0]
        section.page_width    = Mm(A4_W_MM)
        section.page_height   = Mm(A4_H_MM)
        section.top_margin    = Mm(MARGIN_MM)
        section.bottom_margin = Mm(MARGIN_MM)
        section.left_margin   = Mm(MARGIN_MM)
        section.right_margin  = Mm(MARGIN_MM)

        style = doc.styles["Normal"]
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after  = Pt(0)

        cell_w, cell_h = cell_dims(cols, rows)
        per_page       = cols * rows
        total          = len(image_paths)
        num_pages      = math.ceil(total / per_page)

        done_count = 0

        for page_idx in range(num_pages):
            if page_idx > 0:
                doc.add_page_break()

            chunk = image_paths[page_idx * per_page : (page_idx + 1) * per_page]

            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table)

            for col_obj in table.columns:
                col_obj.width = Mm(cell_w)
            for row_obj in table.rows:
                set_row_height(row_obj, cell_h)
                set_row_cant_split(row_obj)

            for i, img_path in enumerate(chunk):
                r_idx = i // cols
                c_idx = i % cols
                cell  = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, 0, 0, 0, 0)

                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run  = para.add_run()

                if os.path.isfile(img_path):
                    buf = resize_for_embed(img_path, cell_w, cell_h)
                    if buf:
                        run.add_picture(buf, width=Mm(cell_w), height=Mm(cell_h))
                        buf.close()
                    else:
                        run.add_picture(img_path, width=Mm(cell_w), height=Mm(cell_h))

                done_count += 1
                if progress_cb:
                    pct = int(done_count / total * 85)
                    progress_cb(pct)

            # fill empty trailing cells
            for i in range(len(chunk), per_page):
                r_idx = i // cols
                c_idx = i % cols
                cell  = table.cell(r_idx, c_idx)
                set_cell_margins(cell, 0, 0, 0, 0)

        if progress_cb:
            progress_cb(90)

        doc.save(output_path)

        if progress_cb:
            progress_cb(100)
        if done_cb:
            done_cb(None)

    except Exception as exc:
        if done_cb:
            done_cb(str(exc))


# ─────────────────────────────────────────────────────────────────────────────
#  Draggable image card widget  — compact: thumbnail + number only
# ─────────────────────────────────────────────────────────────────────────────

class ImageCard(tk.Frame):
    CARD_W = 126        # fixed card width  (px)
    CARD_H = 148        # fixed card height (px)
    DRAG_THRESHOLD = 6  # px before a press becomes a drag

    def __init__(self, master, path: str, index: int,
                 on_drag_start=None, on_drag_move=None, on_drag_end=None,
                 on_delete=None, **kw):
        super().__init__(master, relief="raised", bd=1,
                         bg="#2b2b2b", cursor="hand2",
                         width=self.CARD_W, height=self.CARD_H, **kw)
        self.pack_propagate(False)   # keep fixed card size

        self.path          = path
        self.index         = index
        self.on_drag_start = on_drag_start
        self.on_drag_move  = on_drag_move
        self.on_drag_end   = on_drag_end
        self.on_delete     = on_delete

        self._drag_start_x  = 0
        self._drag_start_y  = 0
        self._grab_offset_x = 0
        self._grab_offset_y = 0
        self._dragging      = False

        # ── thumbnail ─────────────────────────────────────────────────────────
        self.thumb = make_thumbnail(path)
        if self.thumb:
            img_lbl = tk.Label(self, image=self.thumb, bg="#2b2b2b")
            img_lbl.pack(padx=3, pady=(8, 2))
        else:
            img_lbl = tk.Label(self, text="🖼", font=("Arial", 32),
                               bg="#2b2b2b", fg="#aaaaaa")
            img_lbl.pack(padx=4, pady=(10, 2))

        # ── number label ──────────────────────────────────────────────────────
        self.num_lbl = tk.Label(self, text=f"#{index + 1}",
                                font=("Arial", 10, "bold"),
                                bg="#2b2b2b", fg="#f0c040")
        self.num_lbl.pack(pady=(0, 6))

        # ── small delete button overlaid in top-right corner ──────────────────
        del_btn = tk.Button(self, text="✕", font=("Arial", 7, "bold"),
                            bg="#c0392b", fg="white", relief="flat",
                            padx=3, pady=1, cursor="hand2",
                            activebackground="#e74c3c", activeforeground="white",
                            command=self._on_delete_click)
        del_btn.place(relx=1.0, x=-2, y=2, anchor="ne")

        # ── drag bindings on card body, image, and number ─────────────────────
        for w in (self, img_lbl, self.num_lbl):
            w.bind("<ButtonPress-1>",   self._on_press)
            w.bind("<B1-Motion>",       self._on_motion)
            w.bind("<ButtonRelease-1>", self._on_release)

    # ── event handlers ────────────────────────────────────────────────────────

    def _on_delete_click(self):
        if self.on_delete:
            self.on_delete(self)

    def _on_press(self, e):
        self._drag_start_x  = e.x_root
        self._drag_start_y  = e.y_root
        self._grab_offset_x = e.x
        self._grab_offset_y = e.y
        self._dragging      = False
        self.config(relief="sunken")

    def _on_motion(self, e):
        if not self._dragging:
            dx = abs(e.x_root - self._drag_start_x)
            dy = abs(e.y_root - self._drag_start_y)
            if max(dx, dy) < self.DRAG_THRESHOLD:
                return
            self._dragging = True
            self.config(relief="raised", bg="#3d6e91")
            if self.on_drag_start:
                self.on_drag_start(self, self._grab_offset_x, self._grab_offset_y)
        if self.on_drag_move:
            self.on_drag_move(self, e.x_root, e.y_root)

    def _on_release(self, e):
        self.config(relief="raised", bg="#2b2b2b")
        was_dragging  = self._dragging
        self._dragging = False
        if was_dragging and self.on_drag_end:
            self.on_drag_end(self, e.x_root, e.y_root)

    def set_index(self, idx: int):
        self.index = idx
        self.num_lbl.config(text=f"#{idx + 1}")


# ─────────────────────────────────────────────────────────────────────────────
#  Main Application
# ─────────────────────────────────────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("A4 Photo Grid → DOCX Generator")
        self.configure(bg="#1e1e1e")
        self.minsize(420, 500)
        self.resizable(True, True)

        self.image_paths: list = []
        self.cards:       list = []

        # drag state
        self._drag_card          = None
        self._drag_orig_index    = None
        self._drag_target_index  = None
        self._drag_grab_offset_x = 0
        self._drag_grab_offset_y = 0
        self._drag_placeholder   = None

        self._build_ui()
        self._check_deps()

    def _check_deps(self):
        missing = []
        if not HAS_PIL:   missing.append("Pillow")
        if not HAS_DOCX:  missing.append("python-docx")
        if missing:
            self.status_var.set(f"⚠ Missing: {', '.join(missing)} — click Install")
            self.install_btn.config(state="normal")

    # ── UI layout ─────────────────────────────────────────────────────────────
    def _build_ui(self):
        # ── toolbar ──────────────────────────────────────────────────────────
        toolbar = tk.Frame(self, bg="#252525", pady=6)
        toolbar.pack(fill="x", side="top")

        btn = dict(font=("Arial", 10, "bold"), relief="flat",
                   padx=12, pady=6, cursor="hand2")

        tk.Button(toolbar, text="➕  Add Images",
                  bg="#4a90d9", fg="white",
                  command=self._add_images, **btn).pack(side="left", padx=6)

        tk.Button(toolbar, text="🗑  Clear All",
                  bg="#c0392b", fg="white",
                  command=self._clear_all, **btn).pack(side="left", padx=2)

        tk.Button(toolbar, text="📄  Generate DOCX",
                  bg="#27ae60", fg="white",
                  command=self._start_generate, **btn).pack(side="right", padx=6)

        self.install_btn = tk.Button(
            toolbar, text="⬇ Install deps",
            bg="#e67e22", fg="white",
            command=self._install_deps,
            state="disabled", **btn)
        self.install_btn.pack(side="right", padx=2)

        # ── layout selector ──────────────────────────────────────────────────
        layout_bar = tk.Frame(self, bg="#1e1e1e", pady=5)
        layout_bar.pack(fill="x", padx=10)

        tk.Label(layout_bar, text="Layout:", bg="#1e1e1e",
                 fg="#aaaaaa", font=("Arial", 10, "bold")).pack(side="left")

        self.layout_var = tk.StringVar(value=DEFAULT_LAYOUT)
        layout_menu = ttk.Combobox(
            layout_bar,
            textvariable=self.layout_var,
            values=list(LAYOUTS.keys()),
            state="readonly",
            font=("Arial", 10),
            width=22)
        layout_menu.pack(side="left", padx=8)

        self.layout_hint = tk.Label(layout_bar, text="",
                                    bg="#1e1e1e", fg="#888888",
                                    font=("Arial", 9, "italic"))
        self.layout_hint.pack(side="left", padx=4)
        self._update_layout_hint()
        layout_menu.bind("<<ComboboxSelected>>",
                         lambda e: (self._refresh_count(),
                                    self._update_layout_hint()))

        # ── paste-path bar ───────────────────────────────────────────────────
        path_bar = tk.Frame(self, bg="#1e1e1e", pady=2)
        path_bar.pack(fill="x", padx=10)

        tk.Label(path_bar, text="Paste path:", bg="#1e1e1e",
                 fg="#aaaaaa", font=("Arial", 9)).pack(side="left")

        self.path_var = tk.StringVar()
        path_entry = tk.Entry(path_bar, textvariable=self.path_var,
                              bg="#333333", fg="white",
                              insertbackground="white",
                              relief="flat", font=("Arial", 10))
        path_entry.pack(side="left", fill="x", expand=True, padx=6, ipady=4)
        path_entry.bind("<Return>", lambda e: self._add_from_entry())

        tk.Button(path_bar, text="Add", bg="#4a90d9", fg="white",
                  relief="flat", padx=8, font=("Arial", 9, "bold"),
                  command=self._add_from_entry).pack(side="left")

        ttk.Separator(self, orient="horizontal").pack(fill="x", pady=4)

        # ── info row ─────────────────────────────────────────────────────────
        info_row = tk.Frame(self, bg="#1e1e1e")
        info_row.pack(fill="x", padx=12)

        self.count_lbl = tk.Label(info_row, text="No images added yet.",
                                  bg="#1e1e1e", fg="#888888",
                                  font=("Arial", 9, "italic"))
        self.count_lbl.pack(side="left")

        tk.Label(info_row, text="Drag cards to reorder",
                 bg="#1e1e1e", fg="#555555",
                 font=("Arial", 8)).pack(side="right")

        # ── scrollable responsive card grid ───────────────────────────────────
        list_frame = tk.Frame(self, bg="#1e1e1e")
        list_frame.pack(fill="both", expand=True, padx=10, pady=4)

        self._canvas = tk.Canvas(list_frame, bg="#1e1e1e", highlightthickness=0)
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical",
                                  command=self._canvas.yview)
        self.card_container = tk.Frame(self._canvas, bg="#1e1e1e")

        self._canvas_window = self._canvas.create_window(
            (0, 0), window=self.card_container, anchor="nw")

        self._canvas.configure(yscrollcommand=scrollbar.set)
        self._canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Resize canvas → stretch container width + re-grid cards
        self._canvas.bind("<Configure>", self._on_canvas_resize)

        # Mousewheel vertical scroll
        self._canvas.bind_all(
            "<MouseWheel>",
            lambda e: self._canvas.yview_scroll(-1 * (e.delta // 120), "units"))

        # ── status / progress bar ─────────────────────────────────────────────
        bottom = tk.Frame(self, bg="#1a1a1a", pady=4)
        bottom.pack(fill="x", side="bottom")

        self.status_var = tk.StringVar(value="Ready")
        tk.Label(bottom, textvariable=self.status_var,
                 bg="#1a1a1a", fg="#aaaaaa",
                 font=("Arial", 9)).pack(side="left", padx=10)

        self.progress = ttk.Progressbar(bottom, length=200, mode="determinate")
        self.progress.pack(side="right", padx=10)

    # ── responsive grid management ────────────────────────────────────────────

    def _on_canvas_resize(self, event=None):
        """Keep card_container as wide as canvas, then re-grid."""
        if event:
            self._canvas.itemconfig(self._canvas_window, width=event.width)
        self.after(10, self._regrid_cards)

    def _current_cols(self):
        """How many card columns fit in the current canvas width."""
        cw = self._canvas.winfo_width() or 500
        slot_w = ImageCard.CARD_W + 8   # card + padx*2
        return max(1, cw // slot_w)

    def _regrid_cards(self):
        """Place all (non-dragged) cards left-to-right, wrapping into rows."""
        cols   = self._current_cols()
        active = [c for c in self.cards if c is not self._drag_card]
        for i, card in enumerate(active):
            row = i // cols
            col = i % cols
            card.grid(row=row, column=col, padx=4, pady=4)
        self.card_container.update_idletasks()
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    # ── image management ──────────────────────────────────────────────────────

    def _current_layout(self):
        return LAYOUTS[self.layout_var.get()]

    def _update_layout_hint(self):
        cols, rows = self._current_layout()
        per = cols * rows
        w, h = cell_dims(cols, rows)
        self.layout_hint.config(
            text=f"{per} image{'s' if per>1 else ''}/page  •  ~{w:.0f}×{h:.0f} mm each")

    def _add_images(self):
        paths = filedialog.askopenfilenames(
            title="Select Images",
            filetypes=[("Image files",
                        "*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.webp"),
                       ("All files", "*.*")])
        for p in paths:
            self._add_path(p)

    def _add_from_entry(self):
        raw = self.path_var.get().strip().strip('"').strip("'")
        if not raw:
            return
        for p in raw.split(";"):
            p = p.strip()
            if p:
                self._add_path(p)
        self.path_var.set("")

    def _add_path(self, path: str):
        path = os.path.normpath(path)
        if not os.path.isfile(path):
            messagebox.showwarning("Not found", f"File not found:\n{path}")
            return
        if path in self.image_paths:
            messagebox.showinfo("Duplicate", f"Already added:\n{path}")
            return
        self.image_paths.append(path)
        self._add_card(path, len(self.image_paths) - 1)
        self._refresh_count()

    def _add_card(self, path: str, idx: int):
        card = ImageCard(self.card_container, path, idx,
                         on_drag_start=self._on_card_drag_start,
                         on_drag_move=self._on_card_drag_move,
                         on_drag_end=self._on_card_drag_end,
                         on_delete=self._delete_card)
        self.cards.append(card)
        self._regrid_cards()

    def _delete_card(self, card):
        idx = self.cards.index(card)
        self.image_paths.pop(idx)
        self.cards.pop(idx)
        card.destroy()
        for i, c in enumerate(self.cards):
            c.set_index(i)
        self._regrid_cards()
        self._refresh_count()

    def _clear_all(self):
        if not self.image_paths:
            return
        if messagebox.askyesno("Clear", "Remove all images?"):
            self.image_paths.clear()
            for c in self.cards:
                c.destroy()
            self.cards.clear()
            self._refresh_count()

    def _refresh_count(self):
        n = len(self.image_paths)
        if n == 0:
            self.count_lbl.config(text="No images added yet.")
            return
        cols, rows = self._current_layout()
        per_page = cols * rows
        pages = math.ceil(n / per_page)
        self.count_lbl.config(
            text=f"{n} image{'s' if n>1 else ''} → "
                 f"{pages} A4 page{'s' if pages>1 else ''} "
                 f"({per_page}/page)")

    # ── drag-and-drop (2-D responsive grid) ───────────────────────────────────

    def _on_card_drag_start(self, card, grab_x, grab_y):
        """Detach card from grid so it can float freely under the cursor."""
        idx = self.cards.index(card)
        self._drag_card          = card
        self._drag_orig_index    = idx
        self._drag_target_index  = idx
        self._drag_grab_offset_x = grab_x
        self._drag_grab_offset_y = grab_y

        # Capture screen position before removing from grid
        # in_= must be the card's direct parent (card_container)
        abs_x = card.winfo_rootx() - self.card_container.winfo_rootx()
        abs_y = card.winfo_rooty() - self.card_container.winfo_rooty()

        # Highlighted placeholder gap at the vacated slot
        self._drag_placeholder = tk.Frame(
            self.card_container,
            bg="#161616",
            width=ImageCard.CARD_W, height=ImageCard.CARD_H,
            highlightbackground="#f0c040", highlightthickness=2)

        card.grid_forget()
        self._relayout_with_placeholder()

        # Float the card above everything, placed relative to its direct parent
        card.lift()
        card.place(in_=self.card_container, x=abs_x, y=abs_y)

    def _on_card_drag_move(self, card, mouse_x_root, mouse_y_root):
        """Move floating card and live-update placeholder position."""
        if self._drag_card is None:
            return

        rel_x = mouse_x_root - self.card_container.winfo_rootx() - self._drag_grab_offset_x
        rel_y = mouse_y_root - self.card_container.winfo_rooty() - self._drag_grab_offset_y
        card.place_configure(x=rel_x, y=rel_y)

        new_idx = self._compute_grid_drop_index(mouse_x_root, mouse_y_root)
        if new_idx != self._drag_target_index:
            self._drag_target_index = new_idx
            self._relayout_with_placeholder()

    def _compute_grid_drop_index(self, mouse_x_root, mouse_y_root):
        """Return insertion index for the current mouse position over the grid."""
        cols   = self._current_cols()
        slot_w = ImageCard.CARD_W + 8
        slot_h = ImageCard.CARD_H + 8

        # Relative to card_container (already accounts for canvas scroll offset)
        cx = self.card_container.winfo_rootx()
        cy = self.card_container.winfo_rooty()
        rel_x = mouse_x_root - cx
        rel_y = mouse_y_root - cy

        col = max(0, min(cols - 1, int(rel_x // slot_w)))
        row = max(0, int(rel_y // slot_h))
        idx = row * cols + col

        others = [c for c in self.cards if c is not self._drag_card]
        return min(idx, len(others))

    def _relayout_with_placeholder(self):
        """Re-grid non-dragged cards, inserting placeholder at target slot."""
        cols   = self._current_cols()
        others = [c for c in self.cards if c is not self._drag_card]

        for c in others:
            c.grid_forget()
        if self._drag_placeholder:
            self._drag_placeholder.grid_forget()

        seq = others[:]
        seq.insert(self._drag_target_index, self._drag_placeholder)

        for i, w in enumerate(seq):
            if w:
                row = i // cols
                col = i % cols
                w.grid(row=row, column=col, padx=4, pady=4)

        self.card_container.update_idletasks()
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    def _on_card_drag_end(self, card, mouse_x_root, mouse_y_root):
        """Commit new order and restore normal grid layout."""
        if self._drag_card is None:
            return

        orig_idx  = self._drag_orig_index
        final_idx = self._drag_target_index

        self.image_paths.insert(final_idx, self.image_paths.pop(orig_idx))
        self.cards.insert(final_idx, self.cards.pop(orig_idx))

        if self._drag_placeholder:
            self._drag_placeholder.destroy()
            self._drag_placeholder = None

        card.place_forget()

        cols = self._current_cols()
        for i, c in enumerate(self.cards):
            c.set_index(i)
            c.grid(row=i // cols, column=i % cols, padx=4, pady=4)

        self._drag_card         = None
        self._drag_orig_index   = None
        self._drag_target_index = None

        self.card_container.update_idletasks()
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    # ── generate DOCX ─────────────────────────────────────────────────────────
    def _start_generate(self):
        if not self.image_paths:
            messagebox.showwarning("No images", "Please add at least one image.")
            return
        if not HAS_DOCX:
            messagebox.showerror("Missing library",
                                 "python-docx is not installed.\n"
                                 "Click 'Install deps' first.")
            return

        out = filedialog.asksaveasfilename(
            title="Save DOCX as…",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")])
        if not out:
            return

        if not out.lower().endswith(".docx"):
            out += ".docx"
        out = os.path.normpath(out)

        cols, rows = self._current_layout()
        self.progress["value"] = 0
        self.status_var.set(f"Building DOCX ({cols}×{rows})… 0%")
        self.update_idletasks()

        paths_copy = list(self.image_paths)
        saved_path = out

        def _run():
            build_docx(
                paths_copy, saved_path,
                cols=cols, rows=rows,
                progress_cb=self._update_progress,
                done_cb=lambda err: self._on_done(err, saved_path))

        threading.Thread(target=_run, daemon=True).start()

    def _update_progress(self, pct: int):
        self.after(0, lambda: self._set_progress(pct))

    def _set_progress(self, pct: int):
        self.progress["value"] = pct
        if pct < 90:
            self.status_var.set(f"Processing images… {pct}%")
        elif pct < 100:
            self.status_var.set(f"Saving DOCX file… {pct}%")
        else:
            self.status_var.set("✅ Done!")
        self.update_idletasks()

    def _on_done(self, err, saved_path=None):
        if err:
            self.after(0, lambda: messagebox.showerror("Error", str(err)))
            self.after(0, lambda: self.status_var.set("❌ Failed"))
        else:
            self.after(0, lambda: self._open_file(saved_path))

    def _open_file(self, path):
        """Show success popup, open the file in Word, highlight in Explorer."""
        import subprocess, sys
        messagebox.showinfo("Done!", f"DOCX generated successfully!\n\n{path}")
        try:
            if sys.platform == "win32":
                os.startfile(path)
                subprocess.Popen(["explorer", "/select,", path])
            elif sys.platform == "darwin":
                subprocess.Popen(["open", path])
            else:
                subprocess.Popen(["xdg-open", path])
        except Exception as e:
            messagebox.showwarning(
                "Could not open file",
                f"File saved but could not open automatically:\n{path}\n\n{e}")

    def _install_deps(self):
        self.status_var.set("Installing dependencies…")
        self.update_idletasks()
        try:
            msg = install_missing()
            messagebox.showinfo("Done",
                                (msg or "All dependencies already installed.") +
                                "\n\nPlease restart the app.")
            self.status_var.set("Restart required.")
        except Exception as e:
            messagebox.showerror("Install failed", str(e))
            self.status_var.set("Install failed.")


# ─────────────────────────────────────────────────────────────────────────────
#  Entry point
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = App()
    app.mainloop()
